"""Render a Document's stored JSON content (sections/blocks) into a real
.docx file, mirroring the same content-to-markup rules the rich editor
uses in `contentToHtmlWithBlocks`/`paragraphChunkToHtml`
(frontend/src/components/DocumentEditorPage.jsx) so the exported file
matches what's shown on screen.
"""
import logging
import re
from io import BytesIO
from pathlib import Path

from django.conf import settings
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

_UL_RE = re.compile(r"^[-*•◦▪]\s+")
_OL_RE = re.compile(r"^\d+[.):]\s+")
_PAGE_BREAK_MARKER = "[[PAGEBREAK]]"
_CONTENT_MARKER_RE = re.compile(r"\[\[(?:BLOCK:([^\]]+)|PAGEBREAK)\]\]")

# Mirrors the paper sizes offered by the Page Layout control in
# DocumentEditorPage.jsx (PAPER_SIZES), in inches instead of px@96dpi.
_PAPER_SIZES_IN = {
    "Letter": (8.5, 11.0),
    "A4": (8.27, 11.69),
    "Legal": (8.5, 14.0),
}
# Matches the editor's own default margins (DocumentEditorPage.jsx pageMarginsInch).
_DEFAULT_MARGINS_IN = {"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25}


def _apply_page_setup(doc: DocxDocument, page_setup: dict) -> float:
    """Apply the page size/margins the user configured in the Page Layout
    control (persisted as document.content['pageSetup']) to the exported
    section, and return the resulting printable width in inches so embedded
    figures can be sized to actually fit inside the margins."""
    width_in, height_in = _PAPER_SIZES_IN.get(
        (page_setup or {}).get("paperSize"), _PAPER_SIZES_IN["Letter"]
    )
    margins = (page_setup or {}).get("margins") or {}

    def _margin(key: str) -> float:
        try:
            value = float(margins.get(key, _DEFAULT_MARGINS_IN[key]))
        except (TypeError, ValueError):
            value = _DEFAULT_MARGINS_IN[key]
        return max(0.0, value)

    top_in, bottom_in, left_in, right_in = _margin("top"), _margin("bottom"), _margin("left"), _margin("right")

    section = doc.sections[0]
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)
    section.top_margin = Inches(top_in)
    section.bottom_margin = Inches(bottom_in)
    section.left_margin = Inches(left_in)
    section.right_margin = Inches(right_in)

    printable_width_in = width_in - left_in - right_in
    if printable_width_in <= 0:
        logger.warning("docx export: margins leave no printable width, falling back to 1in")
        printable_width_in = 1.0
    return printable_width_in


def _is_list_line(line: str) -> bool:
    return bool(_UL_RE.match(line) or _OL_RE.match(line))


def _heading_level(title: str) -> int:
    t = title.strip()
    if re.match(r"^chapter\s+\d", t, re.IGNORECASE) or re.match(r"^chapter\s+[ivxlc]+\b", t, re.IGNORECASE):
        return 1
    if re.match(r"^\d+\.\d+\.\d+", t):
        return 3
    return 2


def _resolve_media_path(src: str) -> Path | None:
    if not src:
        return None
    media_url = settings.MEDIA_URL or "/media/"
    if src.startswith(media_url):
        rel = src[len(media_url):]
    else:
        rel = src.lstrip("/")
        if rel.startswith("media/"):
            rel = rel[len("media/"):]
    path = Path(settings.MEDIA_ROOT) / rel
    return path if path.exists() else None


def _add_table_block(doc: DocxDocument, block: dict) -> None:
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    n_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if n_cols:
        table = doc.add_table(rows=0, cols=n_cols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            table.style = "Table Grid"
        if headers:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                run = cells[i].paragraphs[0].add_run(str(h))
                run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, cell in enumerate(row):
                if i < n_cols:
                    cells[i].text = str(cell)
    caption = block.get("caption")
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
    doc.add_paragraph()


def _add_image_block(doc: DocxDocument, block: dict, printable_width_in: float) -> None:
    path = _resolve_media_path(block.get("src") or "")
    image_added = False
    if path:
        try:
            doc.add_picture(str(path), width=Inches(min(5.5, printable_width_in)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_added = True
        except Exception:
            logger.warning("docx export: failed to embed image %s", path, exc_info=True)
    else:
        logger.warning("docx export: could not resolve media path for image src=%r", block.get("src"))

    if not image_added:
        p = doc.add_paragraph()
        run = p.add_run(
            f"[Figure could not be embedded: {block.get('caption') or block.get('src') or 'unknown'}]"
        )
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = block.get("caption")
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_block(doc: DocxDocument, block: dict | None, printable_width_in: float) -> None:
    if not block:
        return
    if block.get("type") == "table":
        _add_table_block(doc, block)
    else:
        _add_image_block(doc, block, printable_width_in)


def _add_paragraph_chunk(doc: DocxDocument, text: str) -> None:
    if not text:
        return
    for para in re.split(r"\n\n+", text):
        lines = [ln.rstrip() for ln in para.split("\n") if ln.strip()]
        if not lines:
            continue
        list_items = [ln for ln in lines if _is_list_line(ln.lstrip())]
        has_intro = len(lines) > 1 and not _is_list_line(lines[0].lstrip()) and len(list_items) >= len(lines) - 1
        is_list = len(list_items) >= 2 and (len(list_items) == len(lines) or has_intro)
        if is_list:
            if has_intro:
                doc.add_paragraph(lines[0].strip())
                items = lines[1:]
            else:
                items = lines
            ordered_count = sum(1 for ln in items if _OL_RE.match(ln.lstrip()))
            style = "List Number" if ordered_count > len(items) - ordered_count else "List Bullet"
            for ln in items:
                clean = _OL_RE.sub("", _UL_RE.sub("", ln.lstrip())).strip()
                doc.add_paragraph(clean, style=style)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for i, ln in enumerate(lines):
                if i > 0:
                    p.add_run().add_break()
                p.add_run(ln)


def _add_section_content(
    doc: DocxDocument, content: str, blocks: list[dict] | None, printable_width_in: float
) -> None:
    blocks = blocks or []
    blocks_by_id = {b.get("block_id"): b for b in blocks if b.get("block_id")}
    normalized = re.sub(r"<br\s*/?>", "\n", content or "", flags=re.IGNORECASE)
    if "[[BLOCK:" not in normalized and _PAGE_BREAK_MARKER not in normalized:
        _add_paragraph_chunk(doc, normalized)
        return
    last = 0
    placed: set[str] = set()
    for m in _CONTENT_MARKER_RE.finditer(normalized):
        _add_paragraph_chunk(doc, normalized[last:m.start()])
        if m.group(0) == _PAGE_BREAK_MARKER:
            doc.add_page_break()
        else:
            block_id = (m.group(1) or "").strip()
            block = blocks_by_id.get(block_id)
            if block:
                placed.add(block_id)
                _add_block(doc, block, printable_width_in)
        last = m.end()
    _add_paragraph_chunk(doc, normalized[last:])
    for b in blocks:
        block_id = (b.get("block_id") or "").strip()
        if block_id and block_id not in placed:
            _add_block(doc, b, printable_width_in)


def build_docx(document) -> BytesIO:
    """Build a Word-compatible .docx for the given Document instance, returning
    a ready-to-send in-memory buffer (seeked to 0).
    """
    doc = DocxDocument()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    content = document.content or {}
    printable_width_in = _apply_page_setup(doc, content.get("pageSetup") or {})

    doc.add_heading(document.title or "Untitled Document", level=0)

    for section in content.get("sections") or []:
        title = section.get("title")
        if title:
            doc.add_heading(title, level=_heading_level(title))
        _add_section_content(doc, section.get("content", ""), section.get("blocks"), printable_width_in)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
